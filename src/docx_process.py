import docx
from googletrans import Translator


class DocxProcessor:
    def __init__(self, file_path):
        self.file_path = file_path
        self.doc = docx.Document(file_path)
        self.paragraphs = self.doc.paragraphs
        self.target_path = "../data/output"

    def translate_paratraphs(self, output_path, start, end):
        if start is None:
            start = 0
        if end is None:
            end = len(self.paragraphs)

        for i in range(len(self.paragraphs)):
            self.translate_one_paragraph(i)

        self.doc.save(output_path)
        print("Document translation is completed.")

    def translate_one_paragraph(self, idx):
        translator = Translator()
        paragraph = self.paragraphs[idx]
        try:
            inline = paragraph.runs
            for j in range(len(inline)):
                translation = translator.translate(
                    inline[j].text, src="chinese (simplified)", dest="en"
                )
                inline[j].text = translation.text
        except Exception as error:
            print("Error translating paragraph {}, {}".format(idx, error))


# Move to unit test
if __name__ == "__main__":
    dp = DocxProcessor("../data/input/中医药现代化研究.docx")
    dp.translate_paratraphs("../data/output/中医药现代化研究.dest.docx")