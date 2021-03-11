import os
import docx
from docx.shared import Pt
import json
from .logger import translation_logger as logger
from .translator import Translator
from ..configs.config import cfg, get_infile_path, get_outfile_path
from .docx_namespaces import DOCX_NAMESPACES
from lxml import etree
from collections import defaultdict
import traceback

class DocxProcessor:
    # TODO: language code variables should move to db related to project/doc
    def __init__(
        self,
        infile=None,
        engine=None,
        sourceLanguageCode=cfg["translate"]["sourceLanguageCode"],
        targetLanguageCode=cfg["translate"]["targetLanguageCode"],
    ):
        self.outfile_path = get_outfile_path(
            infile, engine, targetLanguageCode=targetLanguageCode
        )
        self.summary_length = cfg["debug"]["summary_length"]
        self.error_count = 0
        self.engine = engine
        self.sourceLanguageCode = sourceLanguageCode
        self.targetLanguageCode = targetLanguageCode
        self.translator = Translator(
            engine=engine,
            sourceLanguageCode=sourceLanguageCode,
            targetLanguageCode=targetLanguageCode,
        )
        # add dict to keep the translated text
        # now:      to minimize the text send to translator
        # future:   to get training data
        self.dict = dict()
        if engine is None:
            self.engine = cfg["translate"]["engine"]
        self.load_doc(infile)

    def load_doc(self, infile):
        logger.info("Loading document...")
        infile_path = get_infile_path(infile)
        self.infile_path = infile_path
        self.doc = docx.Document(infile_path)
        self.paragraphs = self.doc.paragraphs

    def translate_doc(self):
        try:
            self.translate_tables_et()
            self.translate_paragraphs_et()
            self.translate_hyperlink_et()
            self.translate_textbox_et()
            #print(self.outfile_path)
            self.doc.save(self.outfile_path)
            logger.info("Translation completed, \ninput: {}\noutput: {}".format(self.infile_path, self.outfile_path))
            return True
        except Exception as error:
            logger.error("Translation failed: {}".format(error))
            return False

    def translate_paragraphs_et(self):
        logger.info("Started translating paragraphs...")
        for paragraph in self.paragraphs:
            self.translate_element_one_et(paragraph._element)
        logger.info("Finished translating paragraphs...")

    #@deprecated(verssion='1.0', reason='changed')
    def translate_paragraphs(self):
        logger.info("Started translating paragraphs...")
        for paragraph in self.paragraphs:
            self.translate_paragraph_one(paragraph)
        logger.info("Finished translating paragraphs...")

    def translate_element_one_et(self, e_element):
        '''
        e_element is element object from lxml
        This function translate any element that has child RUNs

        using ET approach instead of docx native method because 
            docx native method has trouble looping through table
            with merged cells
        '''
        run = None
        e_text = None
        text = ""
        try:
            # collect all the runs in the element for text
            # if this approcha collects too deep, will try to use 
            #   for child in elem:
            #       if child.tag = RUN
            for child in e_element:
                if child.tag == DOCX_NAMESPACES["RUN"]:
                    t_count = 0
                    for t in child.iter(DOCX_NAMESPACES["TEXT"]):
                        t_count += 1
                        text += t.text
                        e_text = t
                    if t_count > 0:
                        run = child
                        e_element.remove(run)

            text = text.strip()
            if text != '':
                # get the translated text out of dict if exists
                if text in self.dict.keys():
                    text_out = self.dict[text]
                else: 
                    #text_out = "output = " + text
                    text_out  = self.translator.translate(text)
                    self.dict[text] = text_out
                e_text.text = text_out
                e_element.append(run)

        except Exception as error:
            self.error_count += 1
            summary = text[: self.summary_length] + "..."
            logger.error(traceback.format_exc())
            logger.error("{}: {}\n{}".format(self.error_count, error, summary))
    
    def translate_hyperlink_et(self):
        logger.info("Started translating hyperlinks...")
        for hl in self.doc.element.iter(DOCX_NAMESPACES['HYPERLINK']):
            self.translate_element_one_et(hl)
        logger.info("Finished translating hyperlinks...")


    def translate_textbox_et(self):
        logger.info("Started translating textboxes...")
        for txbx in self.doc.element.iter(DOCX_NAMESPACES["TEXTBOX_CONTENT"]):
            for para in txbx.iter(DOCX_NAMESPACES['PARAGRAPH']):
                self.translate_element_one_et(txbx)
        logger.info("Finished translating textboxes...")
                
    #@deprecated(version='1.0', reason='changed')
    def translate_table_one(self, e_table):
        '''
        e_table is element object from lxml
        '''
        for cell in e_table.iter(DOCX_NAMESPACES["CELL"]):
            for para in cell.iter(DOCX_NAMESPACES["PARAGRAPH"]):
                self.translate_element_one_et(para)

    def translate_tables_et(self):
        logger.info("Started translating tables...")
        for table in self.doc.tables:
            self.translate_table_one(table._element)
        logger.info("Finished translating tables...")

    #@deprecated(version='1.0', reason='changed')
    def translate_paragraph_one(self, paragraph):
        try:
            if paragraph.text != "":
                if self.targetLanguageCode == 'en':
                    paragraph.style.font.name = 'Times'
                elif self.targetLanguageCode == 'zh-cn':
                    paragraph.style.font.name = "songti TC"

                paragraph.text = self.translator.translate(paragraph.text)

        except Exception as error:
            self.error_count += 1
            summary = paragraph.text[: self.summary_length] + "..."
            logger.error("{}: {}\n{}".format(self.error_count, error, summary))

    #@deprecated(version='1.0', reason='changed')
    def translate_tables(self):
        logger.info("Started translating tables...")
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.translate_paragraph_one(paragraph)
        logger.info("Finished translating tables...")
