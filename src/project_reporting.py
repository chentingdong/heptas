import re, json, os, hashlib, string
import docx
import csv
from .logger import reporting_logger as logger
from ..configs.config import cfg, get_infile_path, get_outfile_path, get_report_path
from .models.project import DocProject

class DocxReporting:
    def __init__(self, project_name, docs):
        self.project_name = project_name
        self.docs = docs
        self.projects = []
        
    def make_projects(self):
        for doc in self.docs:
            project = self.make_project_one(doc)
            self.projects.append(project)
            
    def make_project_one(self, doc):
        project = DocProject()
        filename = doc['file']
        project.infile_name = filename
        infile_path = get_infile_path(filename)
        docxObj = docx.Document(infile_path)
        project.infile_size = os.path.getsize(infile_path)
        project.page_count = self.count_pages(docxObj)
        project.wordcount_all, project.wordcount_word = self.count_words(docxObj)
        project.md5sum = self.get_md5sum(infile_path)
        
        outfile_path = get_outfile_path(filename, targetLanguageCode=doc['targetLanguageCode'])
        project.outfile_name = os.path.basename(outfile_path)
        project.outfile_size = os.path.getsize(outfile_path)
        project.billing_status = ""
        return project
    
    def get_md5sum(self, filepath):
        content = open(filepath, 'rb').read()
        md5_hash = hashlib.md5()
        md5_hash.update(content)
        digest = md5_hash.hexdigest()
        return digest
    
    def count_words(self, doc):
        all_counts, word_counts = 0, 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                all_chars = run.text
                all_counts += len(all_chars)
                word_chars = re.sub(r'[^\w\s]', '', all_chars) 
                word_counts += len(word_chars)
        return all_counts, word_counts

    def count_pages(self, doc):
        softPages, hardPages = 0, 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if "lastRenderedPageBreak" in run._element.xml:
                    softPages += 1
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    hardPages += 1
        pages = softPages + hardPages
        logger.info('pages: {}, {}'.format(softPages, hardPages))
        return pages

    def report(self):
        self.reportCsv()

    def report_csv(self):
        self.make_projects()
        try:
            reportfile_path = get_report_path(self.project_name)
            with open(reportfile_path, 'w') as csvfile:
                header = self.projects[0].__dict__.keys()
                writer = csv.DictWriter(csvfile, fieldnames = header)
                writer.writeheader()
                for project in self.projects:            
                    data = project.__dict__
                    logger.debug(data)
                    writer.writerow(data)
            logger.info("Reports written to file {}".format(reportfile_path))
            return True
        except Exception as error:
            logger.error('Reporting pipeline error, {}'.format(error))