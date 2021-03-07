import docx




filepath = '../data/input/2.6.7 毒理学试验列表-LY03009-20210122.docx'
filepath = '2.6.7 毒理学试验列表-LY03009-20210122.tbtest.docx' 

docx = docx.Document(filepath)


i = 0 
print("tables = {}".format(len(docx.tables)))
countTable = 0
for table in docx.tables:
    countTable += 1
    print("\trows = {}".format(len(table.rows)))
    rowCount = 0
    for row in table.rows:
        rowCount += 1
        #if rowCount == 5:
        print(row._element.xml)
        print("\t\tcells = {}".format(len(row.cells)))
        cellCount = 0
        for cell in row.cells:
            cellCount += 1
            print("\t\t\tparagrahs = {}".format(len(cell.paragraphs)))
            #print(cell._element.xml)
            for para in cell.paragraphs:
                if para.text != "":
                    print("\t\t\t\tcounter {}, text = \"{}\"".format(i, para.text))
                    i += 1
                if i > 100:
                    break

