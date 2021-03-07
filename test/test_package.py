from docx import Document
from docx.package import Package
from docx.oxml import parse_xml, OxmlElement
from lxml import etree


CONTENT_NAMESPACE = "{http://schemas.openxmlformats.org/package/2006/content-types}"
WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
MARKUP_COMPATIBILITY = '{http://schemas.openxmlformats.org/markup-compatibility/2006}'
SCM = '{urn:schemas-microsoft-com:vml}'
PARAGRAPH = WORD_NAMESPACE + 'p'
TEXT = WORD_NAMESPACE + 't'
STYLE = WORD_NAMESPACE + 'pStyle'
TABLE = WORD_NAMESPACE + 'tbl'
ROW   = WORD_NAMESPACE + 'tr'
CELL  = WORD_NAMESPACE + 'tc'
RUN = WORD_NAMESPACE + 'r'
ALTERNATE_CONTENT = MARKUP_COMPATIBILITY + 'AlternateContent'
TEXTBOX = SCM + 'textbox'
TEXTBOX_CONTENT = WORD_NAMESPACE + 'txbxContent'
HYPERLINK = WORD_NAMESPACE + 'hyperlink'

doc = "../data/input/CMC test 5-SHARK细胞培养工艺描述 30116-R2.docx"
doc_out = "../data/input/CMC test 5-SHARK细胞培养工艺描述 30116-R2.tb.docx"
doc = '../data/input/2.6.7 毒理学试验列表-LY03009-20210122.tbtest.docx'

#p = Package.open(doc)
def test_hyperlink():
    doc = '../data/input/2.6.7 毒理学试验列表-LY03009-20210122.tbtest.docx'
    docx = Document(doc)
    
    element = docx._element
    for table in element.iter(TABLE):
        for cell in table.iter(CELL):
            for para in cell.iter(PARAGRAPH):
                for child in para:
                    print(child.tag)
        break

def test_math():
    doc = '../data/input/CMC test 4_SHARK 自由巯基报告.math.docx'
    docx = Document(doc)
    counter = 0

    for para in docx.paragraphs:
        counter += 1
        if counter == 3:
            ele = para._element
            for child in ele.iter():
                print(child.tag)
            break

def test_write_para():
    ''' 
    get the text from paragraph
    replace it with new text
    save it back to doc
    '''
    doc = "../data/input/2.6.7 毒理学试验列表-LY03009-20210122.docx"
    doc_out = "../data/input/2.6.7 毒理学试验列表-LY03009-20210122.tbtest.docx"

    docx = Document(doc)
    for table in docx.tables:
        tab = table._element
        for cell in tab.iter(CELL):
            run = None
            text = ""
            for para in cell.iter(PARAGRAPH):
                for run in para.iter(RUN):
                    for txt in run.iter(TEXT):
                        text += txt.text
                    para.remove(run)
                if not run is None:
                    txt = run.find(TEXT)
                    txt.text = "TEST" + text + "TEST"
            if not run is None:
                para.append(run)
    docx.save(doc_out)

# different files in word package
#for part in p.iter_parts():
#    print(part.partname + "\n");

# explore Document
def test_textbox():
    docx = Document(doc)
    counter = 0
    for para in docx.paragraphs:
        counter += 1
        element = para._element
        #print(element.tag)
        for name in element.iter(TEXTBOX_CONTENT):
            for p in name.iter(PARAGRAPH):
                for r in p.iter(CONTENT):
                    for t in r.iter(TEXT):
                        #print(t.text)
                        t.text = 'TTTESTTTT'
            #print(name.xpath("//text()"))
            #print(etree.tostring(name, method='TEXT'));
    docx.save(doc_out);

def test_paragraph():
    docx = Document(doc)
    for para in docx.paragraphs:
        counter += 1
        if counter == 3:
            #print(para._element.xml)
            element = para._element
            print(len(element))
            print(element.tag)
            #print(element.find(CONTENT))
            if element.find(CONTENT):
                for child in element:
                    print(child.tag)
                    if child.find(ALTERNATE_CONTENT):
                        for c in child:
                            print(c.tag)
                    #print(child.xml)
                    #print(etree.tostring(child, method='text'))
                    #print(child.xpath("//text()"))
            #element.find("<w:r>")
            #print(para._element)
            #break
    #print(parse_xml(docx.l))
    
    #print(docx.part.blob)

#test_write_para()
#test_hyperlink()
test_math()

